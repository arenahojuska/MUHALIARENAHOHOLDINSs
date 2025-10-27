import subprocess 
import os
import webbrowser
from datetime import datetime
from robot.api import ExecutionResult, ResultVisitor
from docx import Document
from docx.shared import Pt
import matplotlib.pyplot as plt
from docx.shared import Inches

# === CONFIGURATION ===
existing_doc_path = r"C:\Users\arenaho.muhali\PycharmProjects\MUHALI_ARENAHO_HOLDINGS\Web_Automation\MUHALI HOLDINGS DAILY STATUS REPORT.docx"

# Step 1: Run Robot Framework Test
robot_command = [
    "robot",
    "--outputdir", "log",
    "main.robot"
]
subprocess.run(robot_command)

# Step 2: Parse output.xml
output_xml_path = os.path.abspath("log/output.xml")
result = ExecutionResult(output_xml_path)
result.suite.visit(ResultVisitor())

# Step 3: Collect test results
test_data = []

def collect_test_results(suite):
    for test in suite.tests:
        test_data.append({
            "name": test.name,
            "status": test.status,
            "message": test.message,
            "elapsed": f"{test.elapsedtime / 1000:.2f} sec"
        })
    for subsuite in suite.suites:
        collect_test_results(subsuite)

collect_test_results(result.suite)

# Step 4: Load existing Word document
doc = Document(existing_doc_path)

# Optional: Add a title before the new section
doc.add_heading("Robot Framework Test Summary", level=1)

# Step 5: Append test results
for i, test in enumerate(test_data, start=1):
    doc.add_heading(f"{i}. {test['name']}", level=2)
    doc.add_paragraph(f"Status  : {test['status']}")
    doc.add_paragraph(f"Elapsed : {test['elapsed']}")
    doc.add_paragraph(f"Message : {test['message']}")
    doc.add_paragraph("")  # sp


pass_count = sum(1 for test in test_data if test['status'] == 'PASS')
fail_count = sum(1 for test in test_data if test['status'] == 'FAIL')

# Labels and data
labels = ['Passed', 'Failed']
sizes = [pass_count, fail_count]
colors = ['#4CAF50', '#F44336']

# Create and save chart
plt.figure(figsize=(5, 5))
plt.pie(sizes, labels=labels, colors=colors, autopct='%1.1f%%', startangle=90)
plt.title("Test Result Summary")
plt.axis('equal')
chart_path = os.path.abspath("test_summary_pie_chart.png")
plt.savefig(chart_path)
plt.close()


doc.add_heading("Overall Test Results", level=1)
doc.add_picture(chart_path, width=Inches(5.5))
doc.add_paragraph("")  # spacing


today_str = datetime.now().strftime("%Y-%m-%d")
new_filename = f"MUHALI HOLDINGS DAILY STATUS REPORT_{today_str}.docx"
new_file_path = os.path.join(os.path.dirname(existing_doc_path), new_filename)
doc.save(new_file_path)


webbrowser.open(f"file://{new_file_path}")
