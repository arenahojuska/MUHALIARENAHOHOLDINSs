import subprocess 
import os
import webbrowser
from datetime import datetime
from robot.api import ExecutionResult, ResultVisitor
from docx import Document
from docx.shared import Pt
import matplotlib.pyplot as plt
from docx.shared import Inches

# === CONFIGURATION ===
# Updated to your specific path
existing_doc_path = r"C:\Users\ARENAHO JUSKA\PycharmProjects\MUHALIARENAHOHOLDINSs\Web Playwright\DAILY STATUS REPORTING\MUHALI HOLDINGS DAILY STATUS REPORT.docx"

# Ensure the script runs from the directory where Report.py is located
# This ensures it finds 'main.robot' and the 'log' folder correctly
script_dir = os.path.dirname(os.path.abspath(__file__))
os.chdir(script_dir)

# Step 1: Run Robot Framework Test
# Check if main.robot exists before running
if not os.path.exists("main.robot"):
    print(f"ERROR: 'main.robot' not found in {script_dir}")
    print("Please ensure main.robot is in the same folder as this script.")
    exit(1)

robot_command = [
    "robot",
    "--outputdir", "log",
    "main.robot"
]
print("Running Robot Framework tests...")
subprocess.run(robot_command)

# Step 2: Parse output.xml
output_xml_path = os.path.abspath("log/output.xml")
if not os.path.exists(output_xml_path):
    print(f"ERROR: Could not find output.xml at {output_xml_path}")
    exit(1)

result = ExecutionResult(output_xml_path)
result.suite.visit(ResultVisitor())

# Step 3: Collect test results
test_data = []

def collect_test_results(suite):
    for test in suite.tests:
        test_data.append({
            "name": test.name,
            "status": test.status,
            "message": test.message,
            "elapsed": f"{test.elapsedtime / 1000:.2f} sec"
        })
    for subsuite in suite.suites:
        collect_test_results(subsuite)

collect_test_results(result.suite)

# Step 4: Load existing Word document
if not os.path.exists(existing_doc_path):
    print(f"ERROR: Template Word document not found at: {existing_doc_path}")
    exit(1)

doc = Document(existing_doc_path)

# Optional: Add a title before the new section
doc.add_heading("Robot Framework Test Summary", level=1)

# Step 5: Append test results
for i, test in enumerate(test_data, start=1):
    doc.add_heading(f"{i}. {test['name']}", level=2)
    doc.add_paragraph(f"Status  : {test['status']}")
    doc.add_paragraph(f"Elapsed : {test['elapsed']}")
    doc.add_paragraph(f"Message : {test['message']}")
    doc.add_paragraph("")  # spacing

pass_count = sum(1 for test in test_data if test['status'] == 'PASS')
fail_count = sum(1 for test in test_data if test['status'] == 'FAIL')

# Step 6: Create and save chart
labels = ['Passed', 'Failed']
sizes = [pass_count, fail_count]
colors = ['#4CAF50', '#F44336']

plt.figure(figsize=(5, 5))
plt.pie(sizes, labels=labels, colors=colors, autopct='%1.1f%%', startangle=90)
plt.title("Test Result Summary")
plt.axis('equal')
chart_path = os.path.abspath("test_summary_pie_chart.png")
plt.savefig(chart_path)
plt.close()

doc.add_heading("Overall Test Results", level=1)
doc.add_picture(chart_path, width=Inches(5.5))
doc.add_paragraph("")  # spacing

# Step 7: Save the new report
today_str = datetime.now().strftime("%Y-%m-%d")
new_filename = f"MUHALI HOLDINGS DAILY STATUS REPORT_{today_str}.docx"
# Save it in the same folder as the template
new_file_path = os.path.join(os.path.dirname(existing_doc_path), new_filename)

try:
    doc.save(new_file_path)
    print(f"Success! Report saved to: {new_file_path}")
    # Step 8: Open the report automatically
    webbrowser.open(f"file://{new_file_path}")
except Exception as e:
    print(f"Error saving the document: {e}")