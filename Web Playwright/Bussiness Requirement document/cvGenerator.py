import streamlit as st
import fitz  # PyMuPDF
from docx import Document
from groq import Groq
import io

# --- 1. CONFIGURATION ---
GROQ_API_KEY = "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"
client = Groq(api_key=GROQ_API_KEY)
MODEL_ID = "llama-3.3-70b-versatile"

# --- 2. PAGE CONFIG ---
st.set_page_config(page_title="Muhali | CV Rebuilder", page_icon=None, layout="wide")

# --- 3. CUSTOM CSS ---
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=DM+Sans:wght@300;400;500;600&family=DM+Mono:wght@400;500&display=swap');

*, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }

html, body, .stApp {
    background-color: #0D0F12 !important;
    color: #E8E8E8 !important;
    font-family: 'DM Sans', sans-serif !important;
}

#MainMenu, footer, header { visibility: hidden; }

.block-container {
    padding: 2.5rem 3rem 4rem 3rem !important;
    max-width: 1300px !important;
}

/* SIDEBAR */
[data-testid="stSidebar"] {
    background-color: #111318 !important;
    border-right: 1px solid #1E2028 !important;
}
[data-testid="stSidebar"] * {
    color: #A0A4B0 !important;
    font-family: 'DM Sans', sans-serif !important;
}
[data-testid="stSidebar"] hr {
    border-color: #1E2028 !important;
    margin: 1.2rem 0 !important;
}

/* PAGE HEADER */
.page-header {
    border-bottom: 1px solid #1E2028;
    padding-bottom: 2rem;
    margin-bottom: 2.5rem;
}
.page-label {
    font-family: 'DM Mono', monospace;
    font-size: 0.72rem;
    letter-spacing: 0.18em;
    text-transform: uppercase;
    color: #4A90D9;
    margin-bottom: 0.6rem;
}
.page-title {
    font-size: 2.2rem;
    font-weight: 300;
    color: #F0F0F0;
    line-height: 1.2;
    letter-spacing: -0.02em;
}
.page-title strong {
    font-weight: 600;
    color: #FFFFFF;
}
.page-subtitle {
    font-size: 0.95rem;
    color: #606470;
    margin-top: 0.5rem;
}

/* SECTION LABELS */
.section-label {
    font-family: 'DM Mono', monospace;
    font-size: 0.68rem;
    letter-spacing: 0.16em;
    text-transform: uppercase;
    color: #404450;
    margin-bottom: 0.6rem;
    display: block;
}

/* CARD PANELS */
.panel {
    background-color: #111318;
    border: 1px solid #1E2028;
    border-radius: 10px;
    padding: 1.5rem;
}

/* FILE UPLOADER */
[data-testid="stFileUploader"] {
    background-color: #111318 !important;
    border: 1px dashed #2A2D38 !important;
    border-radius: 10px !important;
    padding: 1rem !important;
    transition: border-color 0.2s ease;
}
[data-testid="stFileUploader"]:hover {
    border-color: #4A90D9 !important;
}
[data-testid="stFileUploader"] label {
    color: #606470 !important;
    font-size: 0.88rem !important;
}

/* TEXT AREA */
.stTextArea textarea {
    background-color: #111318 !important;
    border: 1px solid #2A2D38 !important;
    border-radius: 8px !important;
    color: #C0C4D0 !important;
    font-family: 'DM Sans', sans-serif !important;
    font-size: 0.88rem !important;
    resize: vertical !important;
}
.stTextArea textarea:focus {
    border-color: #4A90D9 !important;
    box-shadow: 0 0 0 2px rgba(74,144,217,0.15) !important;
}
.stTextArea label {
    color: #606470 !important;
    font-size: 0.82rem !important;
}

/* GENERATE BUTTON */
.stButton > button {
    width: 100% !important;
    background-color: #4A90D9 !important;
    color: #FFFFFF !important;
    font-family: 'DM Sans', sans-serif !important;
    font-weight: 500 !important;
    font-size: 0.9rem !important;
    letter-spacing: 0.04em !important;
    border: none !important;
    border-radius: 8px !important;
    height: 3rem !important;
    transition: background-color 0.2s ease, transform 0.1s ease, box-shadow 0.2s ease !important;
}
.stButton > button:hover {
    background-color: #3A7EC8 !important;
    transform: translateY(-1px) !important;
    box-shadow: 0 4px 20px rgba(74, 144, 217, 0.25) !important;
}
.stButton > button:active {
    transform: translateY(0px) !important;
}

/* DOWNLOAD BUTTON */
[data-testid="stDownloadButton"] > button {
    background-color: transparent !important;
    color: #4A90D9 !important;
    border: 1px solid #2A3A50 !important;
    border-radius: 8px !important;
    font-family: 'DM Sans', sans-serif !important;
    font-weight: 500 !important;
    font-size: 0.88rem !important;
    height: 3rem !important;
    width: 100% !important;
    transition: all 0.2s ease !important;
}
[data-testid="stDownloadButton"] > button:hover {
    background-color: #4A90D9 !important;
    color: #FFFFFF !important;
    border-color: #4A90D9 !important;
}

/* CV PREVIEW BOX */
.stTextArea [data-baseweb="base-input"] {
    background-color: #0A0C0F !important;
}

/* DIVIDER */
hr {
    border: none !important;
    border-top: 1px solid #1E2028 !important;
    margin: 2rem 0 !important;
}

/* SUCCESS / ERROR */
[data-testid="stAlert"] {
    background-color: #111318 !important;
    border: 1px solid #2A2D38 !important;
    border-radius: 8px !important;
    color: #A0A4B0 !important;
    font-size: 0.88rem !important;
}

/* FOOTER */
.footer-bar {
    margin-top: 4rem;
    padding-top: 1.5rem;
    border-top: 1px solid #1A1D24;
    display: flex;
    justify-content: space-between;
    align-items: center;
}
.footer-text {
    font-size: 0.78rem;
    color: #2E3240;
    font-family: 'DM Mono', monospace;
    letter-spacing: 0.06em;
}
.status-badge {
    display: inline-flex;
    align-items: center;
    gap: 0.4rem;
    font-size: 0.75rem;
    color: #3CB97A;
    font-family: 'DM Mono', monospace;
}
.status-dot {
    width: 6px;
    height: 6px;
    border-radius: 50%;
    background-color: #3CB97A;
    display: inline-block;
}
</style>
""", unsafe_allow_html=True)


# --- 4. HELPERS ---
def extract_text(uploaded_file):
    ext = uploaded_file.name.split('.')[-1].lower()
    try:
        if ext == "pdf":
            uploaded_file.seek(0)
            with fitz.open(stream=uploaded_file.read(), filetype="pdf") as doc:
                return "".join([page.get_text() for page in doc])
        elif ext == "docx":
            doc = Document(uploaded_file)
            return "\n".join([p.text for p in doc.paragraphs])
        return uploaded_file.read().decode("utf-8")
    except Exception as e:
        st.error(f"Error reading file: {e}")
        return None


def create_docx(content: str) -> bytes:
    doc = Document()

    # Set document styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = None

    for line in content.split('\n'):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        # Detect markdown-style headers
        if stripped.startswith("### "):
            p = doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            p = doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        else:
            doc.add_paragraph(stripped)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def call_groq(cv_text: str, job_desc: str, applicant_name: str, target_role: str) -> str:
    rebuild_prompt = f"""You are an elite professional CV writer and career strategist with 15+ years of experience helping candidates land jobs at top companies worldwide.

Your task is to COMPLETELY REWRITE the candidate's CV to be a perfect, tailored match for the specific job they are applying for.

=== CANDIDATE INFORMATION ===
Name: {applicant_name if applicant_name else "As found in the CV"}
Target Role: {target_role if target_role else "As specified in the Job Description"}

=== ORIGINAL CV ===
{cv_text}

=== TARGET JOB DESCRIPTION ===
{job_desc}

=== REWRITING INSTRUCTIONS ===

1. PROFESSIONAL SUMMARY (3-5 sentences):
   - Open with the candidate's years of experience and the EXACT job title from the JD.
   - Highlight 2-3 skills or achievements that directly match the JD's top requirements.
   - End with a value statement showing what the candidate brings to THIS specific company.
   - Mirror the tone and language used in the JD.

2. CORE COMPETENCIES / SKILLS:
   - Extract every technical skill, tool, framework, and methodology mentioned in the JD.
   - Cross-reference with the candidate's existing skills and PRIORITIZE matches.
   - Add any transferable skills from the CV that support the JD requirements.
   - Group skills logically: e.g., Testing Tools | Programming Languages | Methodologies | Soft Skills.

3. PROFESSIONAL EXPERIENCE:
   - Keep all real employers, job titles, and dates — do NOT fabricate facts.
   - Rewrite EVERY bullet point using strong action verbs (Engineered, Automated, Delivered, Optimised, Led, Reduced, Increased).
   - Inject keywords and phrases directly from the JD into bullet points where truthfully applicable.
   - Quantify achievements wherever possible: percentages, time saved, team sizes, coverage metrics.
   - Ensure the most JD-relevant experience is described with the most detail.

4. EDUCATION:
   - Keep factual. List degree, institution, and graduation year as-is.
   - Add relevant coursework or academic projects only if they support the JD requirements.

5. CERTIFICATIONS / ADDITIONAL SECTIONS:
   - Retain all real certifications from the CV.
   - Suggest (in brackets) any certifications that would strengthen this application based on the JD.

6. ATS OPTIMISATION:
   - Naturally embed the top 10 keywords from the JD throughout the CV.
   - Use the EXACT terminology from the JD (e.g., if JD says "Robot Framework", do not write "Robotframework").
   - Avoid tables, graphics, or columns that ATS systems cannot parse.

7. FORMATTING:
   - Use clean section headers: PROFESSIONAL SUMMARY | CORE COMPETENCIES | PROFESSIONAL EXPERIENCE | EDUCATION | CERTIFICATIONS
   - Use bullet points (- ) under experience sections.
   - Keep the CV to 2 pages worth of content.
   - Do NOT include a photo placeholder, references section, or "References available on request".

OUTPUT: Return ONLY the fully rewritten CV. No explanations, no commentary, no preamble. Start directly with the candidate's name and contact details.
"""

    response = client.chat.completions.create(
        model=MODEL_ID,
        messages=[
            {
                "role": "system",
                "content": "You are an elite CV writer. You return only the fully rewritten CV document with no extra commentary, explanations, or preamble."
            },
            {
                "role": "user",
                "content": rebuild_prompt
            }
        ],
        max_tokens=4096,
        temperature=0.4
    )

    return response.choices[0].message.content


# --- 5. SIDEBAR ---
with st.sidebar:
    st.markdown('<div style="font-size:1.05rem;font-weight:600;color:#FFFFFF;letter-spacing:0.04em;text-transform:uppercase;">Muhali Framework</div>', unsafe_allow_html=True)
    st.markdown("---")
    st.markdown('<div style="font-size:0.82rem;color:#606470;line-height:1.8;">Module<br><span style="color:#A0A4B0;font-weight:500;">CV Rebuilder</span></div>', unsafe_allow_html=True)
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown('<div style="font-size:0.82rem;color:#606470;line-height:1.8;">AI Engine<br><span style="color:#A0A4B0;font-weight:500;">Meta Llama 3.3 via Groq</span></div>', unsafe_allow_html=True)
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown('<div style="font-size:0.82rem;color:#606470;line-height:1.8;">Accepted CV Formats<br><span style="color:#A0A4B0;font-weight:500;">PDF, DOCX, TXT</span></div>', unsafe_allow_html=True)
    st.markdown("---")
    st.markdown('<div style="font-size:0.8rem;color:#404450;line-height:1.6;">Upload your existing CV and paste any job description. The engine will rewrite your CV to be a precise match — ATS optimised and role-specific.</div>', unsafe_allow_html=True)


# --- 6. MAIN CONTENT ---
st.markdown("""
<div class="page-header">
    <div class="page-label">Career Intelligence Suite</div>
    <div class="page-title">CV <strong>Rebuilder</strong></div>
    <div class="page-subtitle">Rewrite your CV for any role — ATS optimised, keyword-matched, and tailored to the job description</div>
</div>
""", unsafe_allow_html=True)

# Optional customisation inputs
st.markdown('<span class="section-label">Step 01 — Customise (Optional)</span>', unsafe_allow_html=True)
opt_col1, opt_col2 = st.columns(2, gap="medium")
with opt_col1:
    applicant_name = st.text_input(
        "Your Full Name",
        placeholder="e.g. Muhali Arenaho Juska",
        label_visibility="visible"
    )
with opt_col2:
    target_role = st.text_input(
        "Target Job Title",
        placeholder="e.g. Senior QA Engineer, Data Analyst, Software Developer",
        label_visibility="visible"
    )

st.markdown("<br>", unsafe_allow_html=True)

# Upload + JD columns
col_a, col_b = st.columns(2, gap="medium")

with col_a:
    st.markdown('<span class="section-label">Step 02 — Upload Your Current CV</span>', unsafe_allow_html=True)
    cv_file = st.file_uploader("Upload CV", type=["pdf", "docx", "txt"], label_visibility="collapsed")

with col_b:
    st.markdown('<span class="section-label">Step 03 — Paste the Job Description</span>', unsafe_allow_html=True)
    job_desc = st.text_area(
        "Job Description",
        height=220,
        placeholder="Paste the full job description here — include responsibilities, requirements, and desired skills...",
        label_visibility="collapsed"
    )

st.markdown("<br>", unsafe_allow_html=True)

# Generate button
if cv_file and job_desc:
    generate_btn = st.button("Rebuild My CV for This Role")

    if generate_btn:
        st.markdown("<br>", unsafe_allow_html=True)
        with st.spinner("Analysing job requirements and rebuilding your CV..."):
            cv_text = extract_text(cv_file)

            if cv_text:
                try:
                    full_cv = call_groq(cv_text, job_desc, applicant_name, target_role)

                    st.markdown("---")
                    st.markdown('<span class="section-label">Step 04 — Preview Rebuilt CV</span>', unsafe_allow_html=True)
                    st.text_area("CV Preview", full_cv, height=520, label_visibility="collapsed")

                    st.markdown("---")
                    st.markdown('<span class="section-label">Step 05 — Export</span>', unsafe_allow_html=True)

                    docx_bytes = create_docx(full_cv)
                    fname = f"CV_{applicant_name.replace(' ', '_') if applicant_name else 'Rebuilt'}.docx"

                    st.download_button(
                        label="Download Rebuilt CV (.docx)",
                        data=docx_bytes,
                        file_name=fname,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )

                except Exception as e:
                    st.error(f"Error: {e}")
elif cv_file and not job_desc:
    st.info("Paste a job description to continue.")
elif job_desc and not cv_file:
    st.info("Upload your CV to continue.")

# Footer
st.markdown("""
<div class="footer-bar">
    <div class="footer-text">MUHALI SOFTWARE SOLUTIONS / 2026</div>
    <div class="status-badge">
        <span class="status-dot"></span>
        GROQ API CONNECTED
    </div>
</div>
""", unsafe_allow_html=True)
