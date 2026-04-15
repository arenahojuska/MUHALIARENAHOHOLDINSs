import streamlit as st
import fitz  # PyMuPDF
from docx import Document
from groq import Groq
import io, os, json, re, tempfile, subprocess

# --- 1. CONFIGURATION ---
GROQ_API_KEY = "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"
client = Groq(api_key=GROQ_API_KEY)
MODEL_ID = "llama-3.3-70b-versatile"

# Path to the Node.js CV builder script (same folder as this file)
CV_BUILDER_JS = os.path.join(os.path.dirname(os.path.abspath(__file__)), "build_cv.js")

# --- 2. PAGE CONFIG ---
st.set_page_config(page_title="Muhali | CV Rebuilder", page_icon=None, layout="wide")

# --- 3. CUSTOM CSS ---
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=DM+Sans:wght@300;400;500;600&family=DM+Mono:wght@400;500&display=swap');

*, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }
html, body, .stApp { background-color: #0D0F12 !important; color: #E8E8E8 !important; font-family: 'DM Sans', sans-serif !important; }
#MainMenu, footer, header { visibility: hidden; }
.block-container { padding: 2.5rem 3rem 4rem 3rem !important; max-width: 1300px !important; }

[data-testid="stSidebar"] { background-color: #111318 !important; border-right: 1px solid #1E2028 !important; }
[data-testid="stSidebar"] * { color: #A0A4B0 !important; font-family: 'DM Sans', sans-serif !important; }
[data-testid="stSidebar"] hr { border-color: #1E2028 !important; margin: 1.2rem 0 !important; }

.page-header { border-bottom: 1px solid #1E2028; padding-bottom: 2rem; margin-bottom: 2.5rem; }
.page-label { font-family: 'DM Mono', monospace; font-size: 0.72rem; letter-spacing: 0.18em; text-transform: uppercase; color: #4A90D9; margin-bottom: 0.6rem; }
.page-title { font-size: 2.2rem; font-weight: 300; color: #F0F0F0; line-height: 1.2; letter-spacing: -0.02em; }
.page-title strong { font-weight: 600; color: #FFFFFF; }
.page-subtitle { font-size: 0.95rem; color: #606470; margin-top: 0.5rem; }

.section-label { font-family: 'DM Mono', monospace; font-size: 0.68rem; letter-spacing: 0.16em; text-transform: uppercase; color: #404450; margin-bottom: 0.6rem; display: block; }

[data-testid="stFileUploader"] { background-color: #111318 !important; border: 1px dashed #2A2D38 !important; border-radius: 10px !important; padding: 1rem !important; transition: border-color 0.2s ease; }
[data-testid="stFileUploader"]:hover { border-color: #4A90D9 !important; }
[data-testid="stFileUploader"] label { color: #606470 !important; font-size: 0.88rem !important; }

.stTextArea textarea { background-color: #111318 !important; border: 1px solid #2A2D38 !important; border-radius: 8px !important; color: #C0C4D0 !important; font-family: 'DM Sans', sans-serif !important; font-size: 0.88rem !important; resize: vertical !important; }
.stTextArea textarea:focus { border-color: #4A90D9 !important; box-shadow: 0 0 0 2px rgba(74,144,217,0.15) !important; }
.stTextArea label { color: #606470 !important; font-size: 0.82rem !important; }
.stTextInput input { background-color: #111318 !important; border: 1px solid #2A2D38 !important; border-radius: 8px !important; color: #C0C4D0 !important; font-family: 'DM Sans', sans-serif !important; }
.stTextInput label { color: #606470 !important; font-size: 0.82rem !important; }

.stButton > button { width: 100% !important; background-color: #4A90D9 !important; color: #FFFFFF !important; font-family: 'DM Sans', sans-serif !important; font-weight: 500 !important; font-size: 0.9rem !important; border: none !important; border-radius: 8px !important; height: 3rem !important; transition: background-color 0.2s ease, transform 0.1s ease !important; }
.stButton > button:hover { background-color: #3A7EC8 !important; transform: translateY(-1px) !important; box-shadow: 0 4px 20px rgba(74,144,217,0.25) !important; }

[data-testid="stDownloadButton"] > button { background-color: transparent !important; color: #4A90D9 !important; border: 1px solid #2A3A50 !important; border-radius: 8px !important; font-family: 'DM Sans', sans-serif !important; font-weight: 500 !important; height: 3rem !important; width: 100% !important; transition: all 0.2s ease !important; }
[data-testid="stDownloadButton"] > button:hover { background-color: #4A90D9 !important; color: #FFFFFF !important; }

hr { border: none !important; border-top: 1px solid #1E2028 !important; margin: 2rem 0 !important; }
[data-testid="stAlert"] { background-color: #111318 !important; border: 1px solid #2A2D38 !important; border-radius: 8px !important; color: #A0A4B0 !important; font-size: 0.88rem !important; }

.footer-bar { margin-top: 4rem; padding-top: 1.5rem; border-top: 1px solid #1A1D24; display: flex; justify-content: space-between; align-items: center; }
.footer-text { font-size: 0.78rem; color: #2E3240; font-family: 'DM Mono', monospace; letter-spacing: 0.06em; }
.status-badge { display: inline-flex; align-items: center; gap: 0.4rem; font-size: 0.75rem; color: #3CB97A; font-family: 'DM Mono', monospace; }
.status-dot { width: 6px; height: 6px; border-radius: 50%; background-color: #3CB97A; display: inline-block; }
</style>
""", unsafe_allow_html=True)


# --- 4. HELPERS ---

def extract_text(uploaded_file):
    ext = uploaded_file.name.split('.')[-1].lower()
    try:
        if ext == "pdf":
            uploaded_file.seek(0)
            with fitz.open(stream=uploaded_file.read(), filetype="pdf") as doc:
                return "".join([page.get_text() for page in doc])
        elif ext == "docx":
            doc = Document(uploaded_file)
            return "\n".join([p.text for p in doc.paragraphs])
        return uploaded_file.read().decode("utf-8")
    except Exception as e:
        st.error(f"Error reading file: {e}")
        return None


def call_groq_json(cv_text: str, job_desc: str, applicant_name: str, target_role: str) -> dict:
    """
    Ask Groq to rewrite the CV and return structured JSON.
    JSON is used by build_cv.js to generate a professional Word document.
    """
    prompt = f"""You are an elite professional CV writer with 15+ years of experience.

Rewrite the candidate's CV to be a perfect tailored match for the job description provided , using a common english  language , dont use  dash , make sure the language being used is human and not robotic 
Return ONLY a valid JSON object — no markdown fences, no explanation, no extra text whatsoever.

JSON structure to return:
{{
  "name": "Full Name",
  "title": "Job Title | Tech Stack keywords matching the JD",
  "contact": {{
    "phone": "phone number from CV",
    "email": "email from CV",
    "github": "github url if present",
    "linkedin": "linkedin url if present",
    "location": "City, Country"
  }},
  "summary": "3-5 sentence professional summary. Open with years of experience + exact JD job title. Match JD tone. End with value statement for this company.",
  "skills": [
    {{ "label": "Category", "items": "Skill1, Skill2, Skill3" }},
    {{ "label": "Category", "items": "Skill1, Skill2, Skill3" }}
  ],
  "experience": [
    {{
      "title": "Job Title",
      "company": "Company Name",
      "location": "City, Country",
      "dates": "Month YYYY – Month YYYY",
      "description": "One sentence describing the role context (optional)",
      "bullets": [
        "Achievement bullet starting with strong action verb, quantified where possible",
        "Achievement bullet"
      ]
    }}
  ],
  "projects": [
    {{
      "name": "Project Name",
      "tech": "Tech Stack used",
      "description": "2-3 sentence description of what was built and achieved"
    }}
  ],
  "education": [
    {{
      "degree": "Degree Name",
      "institution": "University Name",
      "year": "Month YYYY",
      "detail": "Specialisation or honours (optional)"
    }}
  ],
  "certifications": [
    {{ "name": "Certification Name", "issuer": "Issuing Body", "year": "YYYY" }}
  ],
  "keywords": ["Keyword1", "Keyword2", "Keyword3"]
}}

REWRITING RULES:
1. Keep ALL real employers, job titles, dates, education — do NOT fabricate facts.
2. SUMMARY: Open with exact years of experience + exact job title from JD.
3. SKILLS: Extract every tool/technology/framework from the JD. Prioritise matches. Group into 5-6 logical categories (Languages, Backend, Frameworks, DevOps, Cloud, Methodologies).
4. EXPERIENCE BULLETS: Strong action verbs (Engineered, Automated, Delivered, Optimised, Led, Reduced, Increased). Inject JD keywords naturally. Quantify with percentages, team sizes, metrics where possible. 5-7 bullets per role.
5. PROJECTS: Keep all real projects from the CV. Rewrite descriptions to emphasise JD-relevant aspects.
6. KEYWORDS: List the top 20 technologies and skills as individual keywords for the keywords section.
7. ATS: Embed top JD keywords naturally. Use EXACT terminology from the JD.

CANDIDATE NAME: {applicant_name if applicant_name else "extract from CV"}
TARGET ROLE: {target_role if target_role else "extract from JD"}

ORIGINAL CV:
{cv_text}

JOB DESCRIPTION:
{job_desc}
"""

    response = client.chat.completions.create(
        model=MODEL_ID,
        messages=[
            {"role": "system", "content": "You are an elite CV writer. Return ONLY valid JSON. No markdown, no explanation, no text outside the JSON object."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=4096,
        temperature=0.3
    )

    raw = response.choices[0].message.content
    # Strip any accidental markdown fences
    cleaned = re.sub(r"```(?:json)?", "", raw).strip().rstrip("`").strip()
    return json.loads(cleaned)


def build_docx_via_node(cv_data: dict) -> bytes:
    """
    Pass the CV JSON to build_cv.js (Node.js) and return the .docx bytes.
    This produces a professionally styled Word document matching Muhali's CV style.
    """
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        out_path = tmp.name

    try:
        result = subprocess.run(
            ["node", CV_BUILDER_JS, json.dumps(cv_data), out_path],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode != 0:
            raise RuntimeError(result.stderr or result.stdout)

        with open(out_path, "rb") as f:
            return f.read()
    finally:
        if os.path.exists(out_path):
            os.unlink(out_path)


def build_docx_fallback(cv_data: dict) -> bytes:
    """
    Pure python-docx fallback if Node.js is unavailable.
    Matches the same style as the PDF — centred name, ruled sections, bullets.
    """
    from docx import Document as DocxDoc
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = DocxDoc()
    for section in doc.sections:
        section.top_margin    = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin   = Cm(1.8)
        section.right_margin  = Cm(1.8)

    BLACK = RGBColor(0x00, 0x00, 0x00)
    DARK  = RGBColor(0x1A, 0x1A, 0x1A)
    GRAY  = RGBColor(0x55, 0x55, 0x55)
    MID   = RGBColor(0x33, 0x33, 0x33)
    FONT  = "Calibri"

    def add_rule(para):
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single')
        bot.set(qn('w:sz'), '6')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), '000000')
        pBdr.append(bot)
        pPr.append(pBdr)

    def gap(after=4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(after)
        return p

    contact = cv_data.get("contact", {})
    parts = [contact.get("phone",""), contact.get("email","")]
    if contact.get("github"):   parts.append(f"github: {contact['github']}")
    if contact.get("linkedin"): parts.append(f"linkedin: {contact['linkedin']}")
    if contact.get("location"): parts.append(contact["location"])
    contact_str = " • ".join([p for p in parts if p])

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run((cv_data.get("name","")).upper())
    r.bold = True; r.font.size = Pt(19); r.font.color.rgb = BLACK; r.font.name = FONT

    # Title
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(3)
    r2 = p2.add_run(cv_data.get("title",""))
    r2.font.size = Pt(10.5); r2.font.color.rgb = DARK; r2.font.name = FONT

    # Contact
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(5)
    r3 = p3.add_run(contact_str)
    r3.font.size = Pt(9); r3.font.color.rgb = MID; r3.font.name = FONT
    add_rule(p3)

    def section_header(title):
        gap(6)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title.upper())
        r.bold = True; r.font.size = Pt(10); r.font.color.rgb = BLACK; r.font.name = FONT
        add_rule(p)
        gap(3)

    def body_para(text, italic=False, color=None, size=9.5):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.font.size = Pt(size); r.font.name = FONT
        r.font.color.rgb = color or DARK
        r.italic = italic
        return p

    def bullet_para(text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(1.5)
        p.paragraph_format.left_indent = Cm(0.5)
        r = p.add_run(text)
        r.font.size = Pt(9.5); r.font.name = FONT; r.font.color.rgb = DARK

    # Summary
    if cv_data.get("summary"):
        section_header("Professional Summary")
        body_para(cv_data["summary"])

    # Skills
    if cv_data.get("skills"):
        section_header("Technical Skills")
        for s in cv_data["skills"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r1 = p.add_run((s.get("label","")) + ": ")
            r1.bold = True; r1.font.size = Pt(9.5); r1.font.name = FONT; r1.font.color.rgb = BLACK
            r2 = p.add_run(s.get("items",""))
            r2.font.size = Pt(9.5); r2.font.name = FONT; r2.font.color.rgb = DARK

    # Experience
    if cv_data.get("experience"):
        section_header("Professional Experience")
        for exp in cv_data["experience"]:
            gap(5)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r1 = p.add_run(f"{exp.get('title','')} | {exp.get('company','')}")
            r1.bold = True; r1.font.size = Pt(9.5); r1.font.name = FONT; r1.font.color.rgb = BLACK
            r2 = p.add_run(f" — {exp.get('location','')} | {exp.get('dates','')}")
            r2.font.size = Pt(9.5); r2.font.name = FONT; r2.font.color.rgb = GRAY; r2.italic = True
            if exp.get("description"):
                body_para(exp["description"])
            for b in exp.get("bullets", []):
                bullet_para(b)

    # Projects
    if cv_data.get("projects"):
        section_header("Projects")
        for proj in cv_data["projects"]:
            gap(4)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r1 = p.add_run(proj.get("name",""))
            r1.bold = True; r1.font.size = Pt(9.5); r1.font.name = FONT; r1.font.color.rgb = BLACK
            if proj.get("tech"):
                r2 = p.add_run(" — " + proj["tech"])
                r2.font.size = Pt(9.5); r2.font.name = FONT; r2.font.color.rgb = GRAY
            body_para(proj.get("description",""))

    # Education
    if cv_data.get("education"):
        section_header("Education")
        for edu in cv_data["education"]:
            gap(3)
            body_para(f"{edu.get('degree','')} — {edu.get('institution','')} ({edu.get('year','')})")
            if edu.get("detail"):
                body_para(edu["detail"], italic=True, color=GRAY)

    # Certifications
    if cv_data.get("certifications"):
        section_header("Certifications")
        for cert in cv_data["certifications"]:
            text = cert.get("name","")
            if cert.get("issuer"): text += " — " + cert["issuer"]
            if cert.get("year"):   text += f" ({cert['year']})"
            bullet_para(text)

    # Keywords
    if cv_data.get("keywords"):
        section_header("Keywords & Technologies")
        body_para(" • ".join(cv_data["keywords"]), color=GRAY, size=8.5)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


# --- 5. SIDEBAR ---
with st.sidebar:
    st.markdown('<div style="font-size:1.05rem;font-weight:600;color:#FFFFFF;letter-spacing:0.04em;text-transform:uppercase;">Muhali Framework</div>', unsafe_allow_html=True)
    st.markdown("---")
    st.markdown('<div style="font-size:0.82rem;color:#606470;line-height:1.8;">Module<br><span style="color:#A0A4B0;font-weight:500;">CV Rebuilder</span></div>', unsafe_allow_html=True)
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown('<div style="font-size:0.82rem;color:#606470;line-height:1.8;">AI Engine<br><span style="color:#A0A4B0;font-weight:500;">Meta Llama 3.3 via Groq</span></div>', unsafe_allow_html=True)
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown('<div style="font-size:0.82rem;color:#606470;line-height:1.8;">Output Format<br><span style="color:#A0A4B0;font-weight:500;">Professional Word Document (.docx)</span></div>', unsafe_allow_html=True)
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown('<div style="font-size:0.82rem;color:#606470;line-height:1.8;">Accepted CV Formats<br><span style="color:#A0A4B0;font-weight:500;">PDF, DOCX, TXT</span></div>', unsafe_allow_html=True)
    st.markdown("---")
    st.markdown('<div style="font-size:0.8rem;color:#404450;line-height:1.6;">Upload your existing CV and paste any job description. The engine will produce a fully tailored, ATS-optimised CV as a professional Word document matching your original CV style.</div>', unsafe_allow_html=True)


# --- 6. MAIN CONTENT ---
st.markdown("""
<div class="page-header">
    <div class="page-label">Career Intelligence Suite</div>
    <div class="page-title">CV <strong>Rebuilder</strong></div>
    <div class="page-subtitle">Rewrite your CV for any role — ATS optimised, keyword-matched, exported as a professional Word document</div>
</div>
""", unsafe_allow_html=True)

st.markdown('<span class="section-label">Step 01 — Customise (Optional)</span>', unsafe_allow_html=True)
opt_col1, opt_col2 = st.columns(2, gap="medium")
with opt_col1:
    applicant_name = st.text_input("Your Full Name", placeholder="e.g. Muhali Arenaho Juska")
with opt_col2:
    target_role = st.text_input("Target Job Title", placeholder="e.g. Senior QA Engineer, Full Stack Developer")

st.markdown("<br>", unsafe_allow_html=True)

col_a, col_b = st.columns(2, gap="medium")
with col_a:
    st.markdown('<span class="section-label">Step 02 — Upload Your Current CV</span>', unsafe_allow_html=True)
    cv_file = st.file_uploader("Upload CV", type=["pdf", "docx", "txt"], label_visibility="collapsed")

with col_b:
    st.markdown('<span class="section-label">Step 03 — Paste the Job Description</span>', unsafe_allow_html=True)
    job_desc = st.text_area(
        "Job Description",
        height=220,
        placeholder="Paste the full job description here — include responsibilities, requirements, and desired skills...",
        label_visibility="collapsed"
    )

st.markdown("<br>", unsafe_allow_html=True)

if cv_file and job_desc:
    if st.button("Rebuild My CV for This Role"):
        st.markdown("<br>", unsafe_allow_html=True)
        with st.spinner("Analysing job requirements and rebuilding your CV..."):
            cv_text = extract_text(cv_file)
            if cv_text:
                try:
                    # Step 1 — Get structured CV data from Groq
                    cv_data = call_groq_json(cv_text, job_desc, applicant_name, target_role)

                    # Step 2 — Build professional Word doc
                    try:
                        # Try Node.js builder first (best quality)
                        docx_bytes = build_docx_via_node(cv_data)
                    except Exception:
                        # Fallback to pure Python if Node unavailable
                        docx_bytes = build_docx_fallback(cv_data)

                    # Step 3 — Show preview + download
                    st.markdown("---")
                    st.markdown('<span class="section-label">Step 04 — CV Ready</span>', unsafe_allow_html=True)

                    name_slug = (applicant_name or cv_data.get("name","Rebuilt")).replace(" ", "_")
                    fname = f"CV_{name_slug}.docx"

                    st.download_button(
                        label="Download Professional CV (.docx)",
                        data=docx_bytes,
                        file_name=fname,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )

                    with st.expander("Show structured CV data (for debugging)"):
                        st.json(cv_data)

                except json.JSONDecodeError as e:
                    st.error(f"Could not parse AI response. Please try again. Detail: {e}")
                except Exception as e:
                    st.error(f"Error: {e}")

elif cv_file and not job_desc:
    st.info("Paste a job description to continue.")
elif job_desc and not cv_file:
    st.info("Upload your CV to continue.")

st.markdown("""
<div class="footer-bar">
    <div class="footer-text">MUHALI SOFTWARE SOLUTIONS / 2026</div>
    <div class="status-badge"><span class="status-dot"></span>GROQ API CONNECTED</div>
</div>
""", unsafe_allow_html=True)
